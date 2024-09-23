import subprocess
import sys
from art import text2art
from docx import Document
import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

# Display ANSI art
def display_banner():
    banner = text2art("c03lh1n0")
    print(banner)

# Run a tool with subprocess and handle errors
def run_command(command):
    try:
        result = subprocess.run(command, shell=True, capture_output=True, text=True)
        return result.stdout
    except Exception as e:
        return f"Error running command: {command}\n{str(e)}"

# Run Nmap Scan
def run_nmap(target):
    print("Running Nmap scan...")
    nmap_cmd = f"nmap -sV {target}"
    return run_command(nmap_cmd)

# Run Nikto Scan
def run_nikto(target):
    print("Running Nikto scan...")
    nikto_cmd = f"nikto -host {target}"
    return run_command(nikto_cmd)

# Run SQLMap Scan
def run_sqlmap(target):
    print("Running SQLMap scan...")
    sqlmap_cmd = f"sqlmap -u {target} --batch"
    return run_command(sqlmap_cmd)

# Run OWASP Zap Scan
def run_owasp_zap(target):
    print("Running OWASP Zap scan...")
    zap_cmd = f"owasp-zap -quick-scan {target}"
    return run_command(zap_cmd)

# Run Dirb Scan
def run_dirb(target):
    print("Running Dirb scan...")
    dirb_cmd = f"dirb {target}"
    return run_command(dirb_cmd)

# Vulnerability severity ratings
def severity_rating(tool, output):
    if "critical" in output.lower():
        return (tool, "High", output)
    elif "warning" in output.lower() or "vulnerable" in output.lower():
        return (tool, "Medium", output)
    else:
        return (tool, "Low", output)

# Generate the docx report
def generate_report(target, results):
    doc = Document()
    doc.add_heading('NexSysHub Pentest Report', 0)
    
    # Adding target and date
    doc.add_paragraph(f'Target: {target}')
    doc.add_paragraph(f'Date: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
    
    # Organize results by severity
    doc.add_heading('Vulnerability Summary by Severity', level=1)
    for severity in ["High", "Medium", "Low"]:
        doc.add_heading(f'{severity} Severity Vulnerabilities:', level=2)
        for tool, sev, result in results:
            if sev == severity:
                doc.add_heading(f'{tool} Results:', level=3)
                doc.add_paragraph(result)
    
    # Save the document
    doc_name = f"NexSysHub_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(doc_name)
    print(f"Report generated: {doc_name}")

# Main function
def main():
    display_banner()
    
    # Target input
    if len(sys.argv) < 2:
        print("Usage: python nexsyshub.py <target_url>")
        sys.exit(1)
    
    target = sys.argv[1]
    print(f"Target: {target}")
    
    # CLI Menu for selecting tools
    print("\nSelect which tools to run:")
    print("1. Run All Tools")
    print("2. Nmap")
    print("3. Nikto")
    print("4. SQLMap")
    print("5. OWASP Zap")
    print("6. Dirb")
    
    choice = input("Enter your choice (1 for all tools): ").strip()
    
    # Tool selection logic
    tools_to_run = []
    if choice == "1" or not choice:
        tools_to_run = [run_nmap, run_nikto, run_sqlmap, run_owasp_zap, run_dirb]
    else:
        if "2" in choice: tools_to_run.append(run_nmap)
        if "3" in choice: tools_to_run.append(run_nikto)
        if "4" in choice: tools_to_run.append(run_sqlmap)
        if "5" in choice: tools_to_run.append(run_owasp_zap)
        if "6" in choice: tools_to_run.append(run_dirb)

    # Parallel execution of tools
    with ThreadPoolExecutor() as executor:
        future_to_tool = {executor.submit(tool, target): tool.__name__ for tool in tools_to_run}
        
        results = []
        for future in as_completed(future_to_tool):
            tool_name = future_to_tool[future]
            try:
                output = future.result()
                results.append(severity_rating(tool_name, output))
            except Exception as exc:
                print(f'{tool_name} generated an exception: {exc}')
    
    # Generate Report
    generate_report(target, results)

if __name__ == "__main__":
    main()
